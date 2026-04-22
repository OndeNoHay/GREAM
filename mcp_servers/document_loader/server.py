"""
document_loader — MCP server para carga y extracción de texto de documentos.

Herramientas expuestas:
  - load_document(path)           → texto extraído (PDF, DOCX, TXT, MD)
  - list_documents(directory)     → rutas de documentos encontrados
  - get_document_metadata(path)   → metadatos (tamaño, páginas, título, autor)

Transporte: stdio (FastMCP)
"""

import json
import pathlib
from typing import Optional

from fastmcp import FastMCP
from mcp.types import ToolAnnotations

mcp = FastMCP("document_loader")

_SUPPORTED_EXTENSIONS = frozenset({".pdf", ".docx", ".txt", ".md"})


# ---------------------------------------------------------------------------
# Internal helpers (importable for unit tests)
# ---------------------------------------------------------------------------

def _extract_pdf(path: pathlib.Path) -> str:
    import pypdf
    reader = pypdf.PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(p for p in pages if p.strip())


def _extract_docx(path: pathlib.Path) -> str:
    from docx import Document
    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _extract_text(path: pathlib.Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


# ---------------------------------------------------------------------------
# MCP tools
# ---------------------------------------------------------------------------

@mcp.tool(annotations=ToolAnnotations(readOnlyHint=True, idempotentHint=True))
def load_document(path: str) -> str:
    """Extract text from a document. Supported formats: PDF, DOCX, TXT, MD."""
    p = pathlib.Path(path)
    if not p.exists():
        return f"Error: file not found: {path}"
    if not p.is_file():
        return f"Error: path is not a file: {path}"

    ext = p.suffix.lower()
    try:
        if ext == ".pdf":
            return _extract_pdf(p)
        if ext == ".docx":
            return _extract_docx(p)
        if ext in {".txt", ".md"}:
            return _extract_text(p)
        supported = ", ".join(sorted(_SUPPORTED_EXTENSIONS))
        return f"Error: unsupported format '{ext}'. Supported: {supported}"
    except Exception as exc:
        return f"Error reading {path}: {exc}"


@mcp.tool(annotations=ToolAnnotations(readOnlyHint=True, idempotentHint=True))
def list_documents(
    directory: str,
    extensions: Optional[list[str]] = None,
    recursive: bool = True,
) -> str:
    """
    List documents in a directory.

    Returns a JSON array of file paths.
    Pass extensions like ['.pdf', '.docx'] to filter; omit for all supported types.
    Set recursive=false to skip subdirectories.
    """
    d = pathlib.Path(directory)
    if not d.exists() or not d.is_dir():
        return json.dumps([])

    exts = frozenset(e.lower() for e in extensions) if extensions else _SUPPORTED_EXTENSIONS
    pattern = "**/*" if recursive else "*"
    files = sorted(
        str(f) for f in d.glob(pattern)
        if f.is_file() and f.suffix.lower() in exts
    )
    return json.dumps(files)


@mcp.tool(annotations=ToolAnnotations(readOnlyHint=True, idempotentHint=True))
def get_document_metadata(path: str) -> str:
    """
    Return metadata for a document as JSON.

    Fields: path, name, extension, size_bytes.
    For PDF: page_count, title, author (when available).
    For DOCX: title, author, paragraph_count.
    """
    p = pathlib.Path(path)
    if not p.exists():
        return json.dumps({"error": f"file not found: {path}"})

    meta: dict = {
        "path": str(p.resolve()),
        "name": p.name,
        "extension": p.suffix.lower(),
        "size_bytes": p.stat().st_size,
    }

    ext = p.suffix.lower()
    try:
        if ext == ".pdf":
            import pypdf
            reader = pypdf.PdfReader(str(p))
            meta["page_count"] = len(reader.pages)
            info = reader.metadata or {}
            if info.get("/Title"):
                meta["title"] = str(info["/Title"])
            if info.get("/Author"):
                meta["author"] = str(info["/Author"])
        elif ext == ".docx":
            from docx import Document
            doc = Document(str(p))
            core = doc.core_properties
            meta["title"] = core.title or ""
            meta["author"] = core.author or ""
            meta["paragraph_count"] = len(doc.paragraphs)
    except Exception as exc:
        meta["metadata_error"] = str(exc)

    return json.dumps(meta)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    mcp.run()
