"""
word_graem MCP server — genera documentos Word (.docx) con estilos ATEXIS.

Herramientas:
  create_document         — crea un .docx desde contenido JSON estructurado
  create_s1000d_changelog — tabla de cambios S1000D estándar
  list_templates          — lista plantillas .dotx disponibles
"""

import json
import pathlib
from typing import Optional

from fastmcp import FastMCP

mcp = FastMCP(name="word_graem")

_PROJECT_ROOT = pathlib.Path(__file__).parent.parent.parent.resolve()
_OUTPUT_DIR = _PROJECT_ROOT / "output"
_TEMPLATES_DIR = _PROJECT_ROOT / "templates"

# ATEXIS brand colors
_DARK_BLUE = (31, 56, 100)   # #1F3864
_LIGHT_BLUE = (46, 117, 182)  # #2E75B6
_ORANGE = (237, 125, 49)     # #ED7D31


def _ensure_output_dir() -> pathlib.Path:
    _OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    return _OUTPUT_DIR


def _apply_heading_style(run, level: int) -> None:
    from docx.shared import RGBColor, Pt
    if level == 1:
        run.font.color.rgb = RGBColor(*_DARK_BLUE)
        run.font.size = Pt(16)
        run.font.bold = True
    elif level == 2:
        run.font.color.rgb = RGBColor(*_LIGHT_BLUE)
        run.font.size = Pt(14)
        run.font.bold = True
    else:
        run.font.color.rgb = RGBColor(*_LIGHT_BLUE)
        run.font.size = Pt(12)
        run.font.bold = True


def _shade_cell(cell, hex_color: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def _build_document(title: str, content: list, template_path: Optional[pathlib.Path] = None):
    from docx import Document
    from docx.shared import RGBColor, Pt

    if template_path and template_path.exists():
        doc = Document(str(template_path))
        for element in list(doc.element.body):
            doc.element.body.remove(element)
    else:
        doc = Document()

    title_para = doc.add_heading(title, level=0)
    for run in title_para.runs:
        run.font.color.rgb = RGBColor(*_DARK_BLUE)
        run.font.size = Pt(18)

    for item in content:
        item_type = item.get("type", "paragraph")

        if item_type == "heading":
            level = int(item.get("level", 1))
            para = doc.add_heading(item.get("text", ""), level=level)
            for run in para.runs:
                _apply_heading_style(run, level)

        elif item_type == "paragraph":
            doc.add_paragraph(item.get("text", ""))

        elif item_type == "table":
            headers = item.get("headers", [])
            rows = item.get("rows", [])
            if not headers:
                continue
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(headers):
                cell = hdr_cells[i]
                cell.text = h
                run = cell.paragraphs[0].runs[0]
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                _shade_cell(cell, "1F3864")
            for row_data in rows:
                row_cells = table.add_row().cells
                for i, cell_text in enumerate(row_data):
                    if i < len(row_cells):
                        row_cells[i].text = str(cell_text)

        elif item_type == "pagebreak":
            doc.add_page_break()

    return doc


@mcp.tool()
def create_document(
    output_filename: str,
    title: str,
    content_json: str,
    template: str = "",
) -> str:
    """
    Crea un documento Word (.docx) con estilos ATEXIS.

    content_json es un array JSON con elementos:
      {"type": "heading", "level": 1, "text": "..."}
      {"type": "paragraph", "text": "..."}
      {"type": "table", "headers": [...], "rows": [[...]]}
      {"type": "pagebreak"}

    template: nombre de fichero .dotx en templates/ (vacío = estilos por defecto).
    Devuelve la ruta absoluta del fichero creado.
    """
    try:
        content = json.loads(content_json) if content_json.strip() else []
    except json.JSONDecodeError as e:
        return f"Error: content_json inválido — {e}"

    if not output_filename.endswith(".docx"):
        output_filename += ".docx"

    safe_name = pathlib.Path(output_filename).name
    if safe_name != output_filename:
        return "Error: output_filename debe ser solo nombre de fichero, sin directorios"

    template_path = None
    if template:
        candidate = (_TEMPLATES_DIR / template).resolve()
        if str(candidate).startswith(str(_TEMPLATES_DIR)) and candidate.exists():
            template_path = candidate

    try:
        doc = _build_document(title, content, template_path)
        out_path = _ensure_output_dir() / safe_name
        doc.save(str(out_path))
        return str(out_path)
    except Exception as e:
        return f"Error al crear documento: {e}"


@mcp.tool()
def create_s1000d_changelog(
    output_filename: str,
    changelog_json: str,
    document_title: str = "Change Record",
    template: str = "",
) -> str:
    """
    Genera una tabla de cambios S1000D estándar como documento Word.

    changelog_json es un array JSON con entradas:
      {"issue": "001", "date": "2024-01-15", "description": "...",
       "author": "JJO", "reason": "..."}

    Devuelve la ruta absoluta del fichero creado.
    """
    try:
        entries = json.loads(changelog_json) if changelog_json.strip() else []
    except json.JSONDecodeError as e:
        return f"Error: changelog_json inválido — {e}"

    if not output_filename.endswith(".docx"):
        output_filename += ".docx"

    safe_name = pathlib.Path(output_filename).name
    if safe_name != output_filename:
        return "Error: output_filename debe ser solo nombre de fichero, sin directorios"

    headers = ["Issue", "Date", "Description", "Author", "Reason"]
    rows = [
        [
            e.get("issue", ""),
            e.get("date", ""),
            e.get("description", ""),
            e.get("author", ""),
            e.get("reason", ""),
        ]
        for e in entries
    ]
    content = [{"type": "table", "headers": headers, "rows": rows}]

    template_path = None
    if template:
        candidate = (_TEMPLATES_DIR / template).resolve()
        if str(candidate).startswith(str(_TEMPLATES_DIR)) and candidate.exists():
            template_path = candidate

    try:
        doc = _build_document(document_title, content, template_path)
        out_path = _ensure_output_dir() / safe_name
        doc.save(str(out_path))
        return str(out_path)
    except Exception as e:
        return f"Error al crear changelog: {e}"


@mcp.tool()
def list_templates() -> str:
    """Lista las plantillas Word (.dotx) disponibles en el directorio templates/."""
    if not _TEMPLATES_DIR.exists():
        return json.dumps({"templates": [], "count": 0})
    templates = [
        {"name": f.name, "size_bytes": f.stat().st_size}
        for f in sorted(_TEMPLATES_DIR.iterdir())
        if f.is_file() and f.suffix.lower() == ".dotx"
    ]
    return json.dumps({"templates": templates, "count": len(templates)})


if __name__ == "__main__":
    mcp.run()
