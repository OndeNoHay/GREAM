"""
Document processor for extracting text from various file formats.

Supports: PDF, DOCX, TXT, Markdown, Excel (XLSX/XLS), XML
"""

import logging
import re
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Optional

from pypdf import PdfReader
from docx import Document as DocxDocument
from openpyxl import load_workbook
import markdown

from app.config import get_settings_manager

logger = logging.getLogger(__name__)


@dataclass
class DocumentChunk:
    """Represents a chunk of processed document."""
    text: str
    page: Optional[int]  # Page number (1-indexed) or None
    chunk_index: int  # Index within the document
    source_file: str  # Original filename
    metadata: dict


@dataclass
class ProcessedDocument:
    """Represents a fully processed document."""
    filename: str
    total_pages: Optional[int]
    chunks: list[DocumentChunk]
    file_type: str


class DocumentProcessor:
    """
    Processes documents and extracts text chunks with metadata.

    Supports multiple file formats and tracks source references
    (filename, page number) for each chunk.
    """

    SUPPORTED_EXTENSIONS = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "docx",
        ".txt": "text",
        ".md": "markdown",
        ".markdown": "markdown",
        ".xlsx": "excel",
        ".xls": "excel",
        ".xml": "xml",
    }

    def __init__(self) -> None:
        """Initialize the document processor."""
        settings_mgr = get_settings_manager()
        ai_settings = settings_mgr.ai_settings
        self.chunk_size = ai_settings.chunk_size
        self.chunk_overlap = ai_settings.chunk_overlap

    def process_file(
        self,
        file_path: Optional[Path] = None,
        file_content: Optional[bytes] = None,
        filename: str = "unknown"
    ) -> ProcessedDocument:
        """
        Process a document file and extract text chunks.

        Args:
            file_path: Path to the file (optional if file_content provided).
            file_content: Raw file bytes (optional if file_path provided).
            filename: Name of the file.

        Returns:
            ProcessedDocument with chunks and metadata.

        Raises:
            ValueError: If file format is not supported.
        """
        # Determine file extension
        if file_path:
            ext = file_path.suffix.lower()
            filename = file_path.name
            with open(file_path, "rb") as f:
                file_content = f.read()
        else:
            ext = Path(filename).suffix.lower()

        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file format: {ext}. "
                f"Supported: {', '.join(self.SUPPORTED_EXTENSIONS.keys())}"
            )

        file_type = self.SUPPORTED_EXTENSIONS[ext]

        logger.info(f"Processing {file_type} file: {filename}")

        # Process based on file type
        if file_type == "pdf":
            return self._process_pdf(file_content, filename)
        elif file_type == "docx":
            return self._process_docx(file_content, filename)
        elif file_type == "text":
            return self._process_text(file_content, filename)
        elif file_type == "markdown":
            return self._process_markdown(file_content, filename)
        elif file_type == "excel":
            return self._process_excel(file_content, filename)
        elif file_type == "xml":
            return self._process_xml(file_content, filename)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def process_text_content(
        self,
        text: str,
        source_name: str = "pasted_text"
    ) -> ProcessedDocument:
        """
        Process raw text content (e.g., from clipboard paste).

        Args:
            text: The text content to process.
            source_name: Name to identify the source.

        Returns:
            ProcessedDocument with chunks.
        """
        logger.info(f"Processing pasted text: {len(text)} chars")

        chunks = self._create_chunks(
            text=text,
            source_file=source_name,
            page=None
        )

        return ProcessedDocument(
            filename=source_name,
            total_pages=None,
            chunks=chunks,
            file_type="text"
        )

    def _process_pdf(
        self,
        content: bytes,
        filename: str
    ) -> ProcessedDocument:
        """Process a PDF file."""
        reader = PdfReader(BytesIO(content))
        total_pages = len(reader.pages)
        all_chunks = []

        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                chunks = self._create_chunks(
                    text=text,
                    source_file=filename,
                    page=page_num,
                    start_index=len(all_chunks)
                )
                all_chunks.extend(chunks)

        logger.info(f"PDF processed: {total_pages} pages, {len(all_chunks)} chunks")

        return ProcessedDocument(
            filename=filename,
            total_pages=total_pages,
            chunks=all_chunks,
            file_type="pdf"
        )

    def _process_docx(
        self,
        content: bytes,
        filename: str
    ) -> ProcessedDocument:
        """Process a DOCX file."""
        doc = DocxDocument(BytesIO(content))

        # Extract all text from paragraphs
        full_text = "\n".join(
            para.text for para in doc.paragraphs if para.text.strip()
        )

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    full_text += "\n" + row_text

        chunks = self._create_chunks(
            text=full_text,
            source_file=filename,
            page=None  # DOCX doesn't have clear page boundaries
        )

        logger.info(f"DOCX processed: {len(chunks)} chunks")

        return ProcessedDocument(
            filename=filename,
            total_pages=None,
            chunks=chunks,
            file_type="docx"
        )

    def _process_text(
        self,
        content: bytes,
        filename: str
    ) -> ProcessedDocument:
        """Process a plain text file."""
        # Try common encodings
        text = None
        for encoding in ["utf-8", "latin-1", "cp1252"]:
            try:
                text = content.decode(encoding)
                break
            except UnicodeDecodeError:
                continue

        if text is None:
            text = content.decode("utf-8", errors="replace")

        chunks = self._create_chunks(
            text=text,
            source_file=filename,
            page=None
        )

        logger.info(f"Text file processed: {len(chunks)} chunks")

        return ProcessedDocument(
            filename=filename,
            total_pages=None,
            chunks=chunks,
            file_type="text"
        )

    def _process_markdown(
        self,
        content: bytes,
        filename: str
    ) -> ProcessedDocument:
        """Process a Markdown file."""
        text = content.decode("utf-8", errors="replace")

        # Convert markdown to plain text (remove formatting)
        # First convert to HTML, then strip tags
        html = markdown.markdown(text)
        plain_text = re.sub(r"<[^>]+>", " ", html)
        plain_text = re.sub(r"\s+", " ", plain_text).strip()

        chunks = self._create_chunks(
            text=plain_text,
            source_file=filename,
            page=None
        )

        logger.info(f"Markdown processed: {len(chunks)} chunks")

        return ProcessedDocument(
            filename=filename,
            total_pages=None,
            chunks=chunks,
            file_type="markdown"
        )

    def _process_xml(
        self,
        content: bytes,
        filename: str
    ) -> ProcessedDocument:
        """Process an XML file by stripping tags and extracting text content."""
        text = content.decode("utf-8", errors="replace")

        # Strip XML/HTML tags and collapse whitespace
        plain_text = re.sub(r"<[^>]+>", " ", text)
        plain_text = re.sub(r"\s+", " ", plain_text).strip()

        chunks = self._create_chunks(
            text=plain_text,
            source_file=filename,
            page=None
        )

        logger.info(f"XML processed: {len(chunks)} chunks")

        return ProcessedDocument(
            filename=filename,
            total_pages=None,
            chunks=chunks,
            file_type="xml"
        )

    def _process_excel(
        self,
        content: bytes,
        filename: str
    ) -> ProcessedDocument:
        """Process an Excel file."""
        workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
        all_chunks = []

        for sheet_num, sheet_name in enumerate(workbook.sheetnames, start=1):
            sheet = workbook[sheet_name]
            rows_text = []

            for row in sheet.iter_rows(values_only=True):
                cell_values = [
                    str(cell) for cell in row
                    if cell is not None and str(cell).strip()
                ]
                if cell_values:
                    rows_text.append(" | ".join(cell_values))

            if rows_text:
                sheet_text = f"[Sheet: {sheet_name}]\n" + "\n".join(rows_text)
                chunks = self._create_chunks(
                    text=sheet_text,
                    source_file=filename,
                    page=sheet_num,  # Use sheet number as "page"
                    start_index=len(all_chunks)
                )
                all_chunks.extend(chunks)

        workbook.close()

        logger.info(
            f"Excel processed: {len(workbook.sheetnames)} sheets, "
            f"{len(all_chunks)} chunks"
        )

        return ProcessedDocument(
            filename=filename,
            total_pages=len(workbook.sheetnames),
            chunks=all_chunks,
            file_type="excel"
        )

    def _create_chunks(
        self,
        text: str,
        source_file: str,
        page: Optional[int],
        start_index: int = 0
    ) -> list[DocumentChunk]:
        """
        Split text into overlapping chunks.

        Args:
            text: The text to split.
            source_file: Original filename.
            page: Page number (if applicable).
            start_index: Starting index for chunk numbering.

        Returns:
            List of DocumentChunk objects.
        """
        # Clean text
        text = re.sub(r"\s+", " ", text).strip()

        if not text:
            return []

        chunks = []
        pos = 0
        chunk_idx = start_index

        while pos < len(text):
            # Get chunk text
            end = pos + self.chunk_size

            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence end within last 20% of chunk
                search_start = pos + int(self.chunk_size * 0.8)
                sentence_end = -1

                for sep in [". ", ".\n", "! ", "? ", "\n\n"]:
                    idx = text.find(sep, search_start, end + 50)
                    if idx != -1:
                        sentence_end = idx + len(sep)
                        break

                if sentence_end != -1:
                    end = sentence_end

            chunk_text = text[pos:end].strip()

            if chunk_text:
                chunks.append(DocumentChunk(
                    text=chunk_text,
                    page=page,
                    chunk_index=chunk_idx,
                    source_file=source_file,
                    metadata={
                        "char_start": pos,
                        "char_end": end,
                    }
                ))
                chunk_idx += 1

            # Move position with overlap
            pos = end - self.chunk_overlap
            if pos <= 0 or end >= len(text):
                break

        return chunks


# Singleton instance
_document_processor: Optional[DocumentProcessor] = None


def get_document_processor() -> DocumentProcessor:
    """Get the document processor instance."""
    global _document_processor
    if _document_processor is None:
        _document_processor = DocumentProcessor()
    return _document_processor
